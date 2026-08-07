"""DOCX parser."""

from docx import Document
from typing import Any, Dict

from .base import BaseParser


class DocxParser(BaseParser):

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    async def parse(self, file_path: str) -> Dict[str, Any]:
        doc = Document(file_path)

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading structure
                style_name = para.style.name if para.style else ""
                if "Heading" in style_name:
                    level = style_name.replace("Heading ", "")
                    paragraphs.append(f"\n## {'#' * int(level) if level.isdigit() else ''} {text}\n")
                else:
                    paragraphs.append(text)

        # Also extract tables
        tables_text = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                rows.insert(1, "---|" * len(rows[0]))
                tables_text.append("<table>\n" + "\n".join(rows) + "\n</table>")

        full_text = "\n\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n" + "\n\n".join(tables_text)

        return {
            "text": full_text,
            "metadata": {},
            "tables": [],
        }
